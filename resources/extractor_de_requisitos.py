#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Extractor de requisitos y reglas de negocio → Markdown (con CONTEXTO GENERAL)
MODIFICADO: Guarda las imágenes procesadas en la carpeta 'images/'.

Entrada: uno o varios ficheros (PDF, DOCX, TXT, PNG/JPG, etc.) o una carpeta
Salida: único fichero Markdown con RF, Reglas, Decisiones, Integraciones, Observaciones y **Contexto General**

Características:
- Texto: extrae y detecta obligaciones/impuestos ("debe", "deberá", "debería", "permitirá", "validará", etc.)
- Tablas: convierte tablas en posibles reglas de negocio (descuentos/tarifas)
- Imágenes/Diagramas: OCR + heurísticas simples con OpenCV para detectar cajas, rombos y flechas
- Resolución de conflictos: resalta discrepancias detectadas
- Trazabilidad: mantiene referencia a archivo y página/índice de imagen
- **Contexto General**: agrega TODO el texto bruto extraído, por cada archivo/origen.
- **Extracción de Imágenes**: Guarda copia de las imágenes analizadas en la carpeta 'images'.

Uso:   
  python req_extractor.py input_path1 [input_path2 ...] -o salida.md --include-context
"""

import argparse
import os
import re
import sys
import json
import io
from dataclasses import dataclass, field
from typing import List, Dict, Tuple, Optional, Union

# Cargas opcionales (con fallback elegante)
try:
    import pdfplumber  # texto PDF (preferido)
except Exception:
    pdfplumber = None

# Fallbacks adicionales para PDF
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    from PyPDF2 import PdfReader  # pypdf/PyPDF2
except Exception:
    PdfReader = None

try:
    import docx  # python-docx
except Exception:
    docx = None

try:
    from PIL import Image
except Exception:
    Image = None

try:
    import pytesseract
except Exception:
    pytesseract = None

try:
    import cv2  # opencv
except Exception:
    cv2 = None

try:
    import camelot  # tablas en PDF
except Exception:
    camelot = None

try:
    import tabula  # tablas en PDF (alternativa)
except Exception:
    tabula = None

try:
    import pandas as pd
except Exception:
    pd = None

try:
    from pdf2image import convert_from_path
except Exception:
    convert_from_path = None

try:
    from rapidfuzz import fuzz
except Exception:
    fuzz = None

# --- Contexto General (almacenamiento global de texto bruto) ---
CONTEXT: Dict[str, List[str]] = {}
USE_OCR: bool = False

# --- Heurísticas para elegir imágenes a OCR ---
OCR_IMG_MIN_SIDE = 64       # mínimo ancho/alto en px
OCR_IMG_MIN_AREA = 25_000   # mínimo área en px
OCR_IMG_MAX_PER_PAGE = 6    # como mucho N imágenes por página (las más grandes)
OUTPUT_IMG_DIR = "images"   # Carpeta de salida para imágenes


def add_context(source: str, text: str) -> None:
    t = (text or '').strip()
    if not t:
        return
    CONTEXT.setdefault(source, []).append(t)


def context_section() -> str:
    md = ["## Contexto General (texto bruto por fuente)\n"]
    if not CONTEXT:
        md.append("_No se extrajo texto de los documentos._\n")
    else:
        def sort_key(src: str):
            # Extraer número de página si existe (e.g. "pag. 12")
            m = re.search(r"pag\.\s*(\d+)", src)
            num = int(m.group(1)) if m else 0
            # Priorizar nombre del archivo y luego número de página
            base = re.sub(r",\s*pag\.\s*\d+", "", src)
            return (base.lower(), num)
        
        for src in sorted(CONTEXT.keys(), key=sort_key):
            combined = "\n\n".join(CONTEXT[src]).strip()
            md.append(f"### Fuente: {src}\n")
            md.append("```\n" + combined + "\n```\n")
    return "\n".join(md)


# -------------------- Utilidades --------------------
# -------------------- Utilidades (PATRONES ACTUALIZADOS SPA+ENG) --------------------

# Combinamos palabras clave de ES y EN para detectar obligaciones
OBLIGATION_PAT = re.compile(
    r"\b("
    # Español
    r"(el\s+sistema\s+)?(debe|deberá|debería|tiene\s+que|deberá\s+permitir|permitirá|validará|calculará|mostrará|registrará|enviará|generará|proporcionará|soportará|incluirá|gestionará)"
    r"|implementación\s+de|capacidad\s+de|requisito\s+de"
    # Inglés
    r"|(the\s+system\s+)?(must|shall|should|will|needs\s+to|has\s+to|is\s+required\s+to|allows|supports|provides|validates|calculates|displays|registers|stores|sends|generates)"
    r"|implementation\s+of|capability\s+to|requirement\s+for"
    r")\b",
    flags=re.IGNORECASE
)

# Patrones condicionales (Si... / If...)
CONDITION_PAT = re.compile(
    r"\b("
    # Español
    r"si|en\s+caso\s+de\s+que|cuando|mientras|a\s+condición\s+de"
    # Inglés
    r"|if|in\s+case|when|while|provided\s+that|unless"
    r")\b", 
    flags=re.IGNORECASE
)

# Patrones de integración
INTEGRATION_PAT = re.compile(
    r"\b("
    # Comunes
    r"SAP|SII|DGT|SFTP|Traffic\s+BI|API[s]?|webhook|JSON|XML"
    # Español
    r"|servicio\s+web|integración\s+con|integrado\s+con|comunicar\s+a|enviar\s+a|conectar\s+con|consumir\s+de"
    # Inglés
    r"|web\s+service|integration\s+with|integrated\s+with|communicate\s+to|send\s+to|connect\s+with|consume\s+from|interface"
    r")\b",
    flags=re.IGNORECASE
)
    
DIAGRAM_TRIGGERS = [
    # Español
    "diagrama de procesos", 
    "flujo de tratamiento", 
    "flujograma", 
    "leyenda",
    # Inglés (NUEVO)
    "process flow",
    "flowchart",
    "diagram",
    "legend",
    "workflow",
    "bpm",
    "swimlane"
]

ARROW_HINTS = ["→", "->", "→", "⇒"]

# Normalización para evitar textos cortados por saltos de línea duros
_HARD_BREAK = re.compile(r"-\n(?=\w)")  # palabras cortadas por guion fin de línea
_SOFT_BREAK = re.compile(r"(?<![\.!?:;])\n(?!\n)")  # salto de línea dentro de párrafo
_MULTISPACE = re.compile(r"[ \t]{2,}")

def normalize_text(s: str) -> str:
    if not s:
        return s
    s = _HARD_BREAK.sub("", s)
    s = _SOFT_BREAK.sub(" ", s)
    s = _MULTISPACE.sub(" ", s)
    # compactar espacios antes de puntuación
    s = re.sub(r"\s+([\.,;:])", r"\1", s)
    return s

def save_debug_image(pil_img: 'Image.Image', source: str) -> None:
    """Guarda la imagen en disco para inspección visual."""
    if pil_img is None:
        return
    try:
        # Asegurar directorio
        os.makedirs(OUTPUT_IMG_DIR, exist_ok=True)
        
        # Limpiar nombre de archivo (eliminar caracteres raros de 'source')
        # source suele ser: "archivo.pdf, pag. 1#diagrama" -> "archivo_pdf__pag__1_diagrama"
        safe_name = re.sub(r"[^\w\-\.]", "_", source)
        
        # Evitar nombres excesivamente largos
        if len(safe_name) > 150:
            safe_name = safe_name[:150]
            
        filename = f"{safe_name}.png"
        path = os.path.join(OUTPUT_IMG_DIR, filename)
        
        pil_img.save(path)
        # print(f"  [IMG] Guardada: {path}") # Descomentar para debug
    except Exception as e:
        print(f"Advertencia: No se pudo guardar imagen de {source}: {e}")

@dataclass
class Item:
    kind: str  # RF | RB | DC | INT | OBS
    text: str
    source: str
    meta: Dict = field(default_factory=dict)

@dataclass
class Extraction:
    rf: List[Item] = field(default_factory=list)
    rb: List[Item] = field(default_factory=list)
    dc: List[Item] = field(default_factory=list)
    it: List[Item] = field(default_factory=list)
    obs: List[Item] = field(default_factory=list)

    def add(self, kind: str, text: str, source: str, meta: Optional[Dict] = None):
        (meta := meta or {})
        item = Item(kind=kind, text=text.strip(), source=source, meta=meta)
        if kind == 'RF':
            self.rf.append(item)
        elif kind == 'RB':
            self.rb.append(item)
        elif kind == 'DC':
            self.dc.append(item)
        elif kind == 'INT':
            self.it.append(item)
        elif kind == 'OBS':
            self.obs.append(item)

    def to_markdown(self, include_context: bool = False) -> str:
        def section(title, items, prefix):
            out = [f"## {title}\n"]
            for i, it in enumerate(items, 1):
                trace = f" _(origen: {it.source})_" if it.source else ""
                out.append(f"- **{prefix}-{i}**: {it.text}{trace}")
            if len(items) == 0:
                out.append("- _Sin hallazgos_.")
            out.append("")
            return "\n".join(out)

        md = [
            "# Especificación derivada (auto-generada)",
            "",
            "> Documento generado automáticamente a partir de insumos suministrados. Revísese manualmente.",
            "",
        ]
        md.append(section("Requerimientos Funcionales (RF)", self.rf, "RF"))
        md.append(section("Reglas de Negocio (RB)", self.rb, "RB"))
        md.append(section("Decisiones / Condiciones (DC)", self.dc, "DC"))
        md.append(section("Integraciones / Flujos (INT)", self.it, "INT"))
        md.append(section("Supuestos y Observaciones", self.obs, "OBS"))
        if include_context:
            md.append(context_section())
        return "\n".join(md)

# -------------------- Parsers --------------------

def read_txt(path: str) -> str:
    with open(path, 'r', encoding='utf-8', errors='ignore') as f:
        return normalize_text(f.read())


def read_pdf_text(path: str) -> Tuple[str, List[Tuple[int, str]]]:
    """Devuelve (texto completo normalizado, [(pagina_idx, texto_pagina_normalizado), ...])
    Orden: pdfplumber → PyMuPDF → PyPDF2
    """
    pages: List[Tuple[int, str]] = []
    all_text = ""

    def _acc(i: int, t: str):
        t = normalize_text(t or "")
        pages.append((i, t))
        return f"\n\n[PAGE {i}]\n" + t

    # 1) pdfplumber
    if pdfplumber is not None:
        try:
            with pdfplumber.open(path) as pdf:
                for i, p in enumerate(pdf.pages, 1):
                    all_text += _acc(i, p.extract_text() or "")
            return all_text.strip(), pages
        except Exception:
            pass  # sigue a fallback

    # 2) PyMuPDF (fitz)
    if fitz is not None:
        try:
            doc = fitz.open(path)
            for i, page in enumerate(doc, 1):
                all_text += _acc(i, page.get_text("text") or "")
            return all_text.strip(), pages
        except Exception:
            pass

    # 3) PyPDF2 / pypdf
    if PdfReader is not None:
        try:
            r = PdfReader(path)
            for i, page in enumerate(r.pages, 1):
                try:
                    t = page.extract_text() or ""
                except Exception:
                    t = ""
                all_text += _acc(i, t)
            return all_text.strip(), pages
        except Exception:
            pass

    return "", []


def read_docx(path: str) -> Tuple[str, List[List[str]]]:
    """Devuelve texto concatenado normalizado y tablas como lista de filas."""
    if docx is None:
        return "", []
    try:
        d = docx.Document(path)
        text = "\n".join([p.text for p in d.paragraphs])
        text = normalize_text(text)
        tables = []
        for tb in d.tables:
            rows = []
            for r in tb.rows:
                rows.append([normalize_text(c.text.strip()) for c in r.cells])
            tables.append(rows)
        return text.strip(), tables
    except Exception:
        return "", []


def extract_pdf_tables(path: str) -> List[Tuple['pd.DataFrame', Optional[int]]]:
    if pd is None:
        return []
    out: List[Tuple['pd.DataFrame', Optional[int]]] = []
    try:
        if camelot is not None:
            tables = camelot.read_pdf(path, pages='all')
            for t in tables:
                page = getattr(t, 'page', None)
                out.append((t.df, int(page) if page else None))
        elif tabula is not None:
            dfs = tabula.read_pdf(path, pages='all', multiple_tables=True)
            for df in dfs:
                out.append((df, None))
    except Exception:
        pass
    return out


def images_from_pdf(path: str) -> List[Tuple[int, 'Image.Image']]:
    if convert_from_path is not None:
        try:
            pil_pages = convert_from_path(path, dpi=200)
            return list(enumerate(pil_pages, 1))
        except Exception as e:
            print(f"[pdf2image] Error al renderizar {path}: {e}")

    if fitz is not None and Image is not None:
        try:
            doc = fitz.open(path)
            result = []
            for i, page in enumerate(doc, 1):
                pix = page.get_pixmap(dpi=150)
                pil = Image.open(io.BytesIO(pix.tobytes("png"))).convert('RGB')
                result.append((i, pil))
            return result
        except Exception as e:
            print(f"[PyMuPDF] Error al renderizar {path}: {e}")

    return []


def extract_pdf_inline_images(path: str) -> List[Tuple[int, int, 'Image.Image', Tuple[int,int]]]:
    results: List[Tuple[int, int, 'Image.Image', Tuple[int,int]]] = []
    if fitz is None or Image is None:
        return results

    import hashlib, io
    try:
        doc = fitz.open(path)
        for pno, page in enumerate(doc, 1):
            raw_imgs = page.get_images(full=True) or []
            kept: List[Tuple[int, int, bytes, int, int]] = []
            seen_hashes = set()

            for img_idx, meta in enumerate(raw_imgs, 1):
                try:
                    xref = meta[0]
                    smask = meta[1]
                    w, h = meta[2], meta[3]
                    if smask: continue
                    if w < OCR_IMG_MIN_SIDE or h < OCR_IMG_MIN_SIDE: continue
                    if (w * h) < OCR_IMG_MIN_AREA: continue

                    pix = fitz.Pixmap(doc, xref)
                    if pix.n >= 5:
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                    png = pix.tobytes("png")

                    hsh = hashlib.sha1(png).hexdigest()
                    if hsh in seen_hashes: continue
                    seen_hashes.add(hsh)

                    kept.append((xref, 0, png, w, h))
                except Exception:
                    continue

            kept.sort(key=lambda t: t[3]*t[4], reverse=True)
            if OCR_IMG_MAX_PER_PAGE is not None and OCR_IMG_MAX_PER_PAGE > 0:
                kept = kept[:OCR_IMG_MAX_PER_PAGE]

            visible_idx = 0
            for xref, _, png, w, h in kept:
                try:
                    pil = Image.open(io.BytesIO(png)).convert("RGB")
                    visible_idx += 1
                    results.append((pno, visible_idx, pil, (w, h)))
                except Exception:
                    continue
    except Exception:
        pass

    return results


def ocr_image(pil_img: 'Image.Image') -> str:
    global USE_OCR
    if not USE_OCR:
        return ""
    if pytesseract is None or Image is None or pil_img is None:
        return ""
    try:
        return normalize_text(pytesseract.image_to_string(pil_img, lang='spa+eng'))
    except Exception:
        return ""


# -------------------- Heurísticas de NLP --------------------

def split_sentences(text: str) -> List[str]:
    parts = re.split(r"(?<=[\.!?])\s+|\n{2,}", text)
    return [p.strip() for p in parts if p and any(ch.isalpha() for ch in p)]


def detect_rf(sent: str) -> bool:
    return bool(OBLIGATION_PAT.search(sent))


def detect_dc(sent: str) -> bool:
    return bool(CONDITION_PAT.search(sent))
    
def detect_int(sent: str) -> bool:
    return bool(INTEGRATION_PAT.search(sent))


def table_to_rules(df, source: str) -> List[Item]:
    items = []
    if pd is None:
        return items
    df2 = df.copy()
    df2.columns = [str(c).strip().lower() for c in df2.columns]
    
    # LISTAS DE PALABRAS CLAVE AMPLIADAS (ES + EN)
    # Columnas de tipo "Condición"
    cond_keywords = [
        "concepto", "tipo", "categoria", "nivel", "tramo", "condicion", "condición", "estado", # ES
        "concept", "type", "category", "level", "tier", "condition", "status" # EN
    ]
    
    # Columnas de tipo "Valor/Resultado"
    val_keywords = [
        "descuento", "%", "porcentaje", "tarifa", "importe", "precio", "monto", # ES
        "discount", "percent", "percentage", "rate", "fee", "amount", "price"   # EN
    ]

    text_cols = [c for c in df2.columns if any(h in c for h in cond_keywords)]
    val_cols = [c for c in df2.columns if any(h in c for h in val_keywords)]

    if text_cols and val_cols:
        for idx, row in df2.iterrows():
            conds = []
            for c in text_cols:
                v = str(row.get(c, '')).strip()
                if v and v.lower() != 'nan':
                    conds.append(f"{c}: '{v}'")
            valores = []
            for c in val_cols:
                v = str(row.get(c, '')).strip()
                if v and v.lower() != 'nan':
                    valores.append(f"{c}: {v}")
            if conds and valores:
                rule = f"Aplicar {', '.join(valores)} si {' y '.join(conds)}."
                items.append(Item(kind='RB', text=rule, source=source, meta={'row': int(idx)}))
        return items

    items.append(Item(kind='RB', text=f"Tabla interpretada como catálogo/parametría. Encabezados: {', '.join(df2.columns[:8])}", source=source))
    return items

# -------------------- Detección de formas en diagramas --------------------

import numpy as np

def analyze_diagram(pil_img: 'Image.Image', source: str) -> List[Item]:
    """
    Analiza un diagrama usando detección de color.
    """
    results = []
    if cv2 is None or Image is None or np is None:
        print(f"Faltan OpenCV o Numpy. Realizando OCR simple sobre {source}.")
        text = ocr_image(pil_img)
        for s in split_sentences(text):
            if detect_dc(s):
                results.append(Item('DC', s, source))
            elif detect_rf(s):
                results.append(Item('RF', s, source))
        return results

    try:
        img_bgr = cv2.cvtColor(np.array(pil_img), cv2.COLOR_RGB2BGR)
        hsv = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2HSV)

        # Azules (RF)
        azul_claro_b = np.array([90, 50, 150])
        azul_claro_a = np.array([115, 255, 255])
        azul_oscuro_b = np.array([100, 150, 50])
        azul_oscuro_a = np.array([130, 255, 200])
        
        # Grises (DC)
        gris_b = np.array([0, 0, 100])
        gris_a = np.array([180, 60, 220]) 

        mascara_azul_claro = cv2.inRange(hsv, azul_claro_b, azul_claro_a)
        mascara_azul_oscuro = cv2.inRange(hsv, azul_oscuro_b, azul_oscuro_a)
        mascara_procesos = cv2.bitwise_or(mascara_azul_claro, mascara_azul_oscuro)
        
        mascara_decisiones = cv2.inRange(hsv, gris_b, gris_a)

        def process_contours(contours, kind_prefix, kind_text):
            for cnt in contours:
                area = cv2.contourArea(cnt)
                if area < 1500:
                    continue
                
                x, y, w, h = cv2.boundingRect(cnt)
                pad = 5
                roi = img_bgr[max(y-pad, 0):min(y+h+pad, img_bgr.shape[0]), 
                              max(x-pad, 0):min(x+w+pad, img_bgr.shape[1])]
                
                try:
                    roi_pil = Image.fromarray(cv2.cvtColor(roi, cv2.COLOR_BGR2RGB))
                except Exception:
                    roi_pil = None

                if roi_pil:
                    # Aquí NO guardamos la ROI por defecto para no saturar,
                    # pero se podría llamar a save_debug_image(roi_pil, source + f"_crop_{x}_{y}")
                    roi_text = ocr_image(roi_pil)
                    content = roi_text.strip()
                    if content:
                        results.append(Item(kind_prefix, f"{kind_text}: {content}", source, 
                                           meta={'bbox':[int(x),int(y),int(w),int(h)]}))

        contours_rf, _ = cv2.findContours(mascara_procesos, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        process_contours(contours_rf, 'RF', 'Proceso')
        
        contours_dc, _ = cv2.findContours(mascara_decisiones, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        process_contours(contours_dc, 'DC', 'Decisión')

    except Exception as e:
        print(f"ERROR en analyze_diagram con {source}: {e}")
        text = ocr_image(pil_img)
        if text:
            results.append(Item('OBS', f"Fallo el análisis de diagrama (color), se usó OCR simple. Texto: {text[:100]}...", source))

    return results

# -------------------- Resolución de conflictos --------------------

def detect_conflicts(items: List[Item]) -> List[Tuple[Item, Item, int]]:
    if fuzz is None:
        return []
    conflicts = []
    for i in range(len(items)):
        for j in range(i+1, len(items)):
            a, b = items[i], items[j]
            sim = fuzz.token_set_ratio(a.text, b.text)
            if sim > 80 and a.text != b.text:
                nums_a = re.findall(r"\d+[\,\.]?\d*", a.text)
                nums_b = re.findall(r"\d+[\,\.]?\d*", b.text)
                if nums_a and nums_b and set(nums_a) != set(nums_b):
                    conflicts.append((a, b, sim))
    return conflicts
    


# -------------------- Pipeline principal --------------------

def process_text(text: str, source: str, extr: Extraction):
    for sent in split_sentences(text):
        if detect_rf(sent):
            extr.add('RF', sent, source)
        if detect_dc(sent):
            extr.add('DC', sent, source)
        if detect_int(sent):
            extr.add('INT', sent, source)
        
        # --- MODIFICACIÓN AQUÍ: AÑADIR KEYWORDS EN INGLÉS ---
        # Palabras clave para detectar Reglas de Negocio (RB) explícitas
        rb_keywords = [
            # Español
            "regla", "restricción", "validación", "fórmula", "porcentaje", 
            # Inglés
            "rule", "restriction", "validation", "formula", "percent", "rate", "calculation"
        ]
        
        # El símbolo % es universal
        if "%" in sent or any(k in sent.lower() for k in rb_keywords):
            extr.add('RB', sent, source)

def process_tablelike(tables: List[Tuple[Union['pd.DataFrame', list], Optional[int]]], source: str, extr: Extraction):
    if pd is None:
        return
    for idx, pair in enumerate(tables, 1):
        try:
            if isinstance(pair, tuple):
                tb, page = pair
            else:
                tb, page = pair, None
            df = tb if isinstance(tb, pd.DataFrame) else pd.DataFrame(tb)
            page_suffix = f", pag. {page}" if page else ""
            rules = table_to_rules(df, f"{source}{page_suffix}#tabla{idx}")
            for it in rules:
                extr.add(it.kind, it.text, it.source, it.meta)
        except Exception:
            continue


def process_image(pil_img: 'Image.Image', source: str, extr: Extraction):
    """
    Procesa una imagen para extraer texto y diagramas.
    MODIFICACIÓN: Guarda la imagen en disco antes de procesar.
    """
    if pil_img is None:
        return

    # --- NUEVO: FILTRO DE TAMAÑO GLOBAL ---
    w, h = pil_img.size
    
    # 1. Filtro por lado mínimo (ej. 64px)
    if w < OCR_IMG_MIN_SIDE or h < OCR_IMG_MIN_SIDE:
        # print(f"Ignorando imagen pequeña en {source}: {w}x{h}")
        return

    # 2. Filtro por área mínima (ej. 25,000px)
    # Nota: Un avatar de 100x100 tiene 10,000px, por lo que sería filtrado aquí.
    if (w * h) < OCR_IMG_MIN_AREA:
        # print(f"Ignorando imagen por área insuficiente en {source}: {w*h}px")
        return
    
    # >>> INICIO MODIFICACIÓN: Guardar imagen
    print(f"Procesando imagen de: {source}")
    save_debug_image(pil_img, source)
    # <<< FIN MODIFICACIÓN

    print(f"Ejecutando OCR sobre imagen extraída de {source}")
    txt = ocr_image(pil_img)
    if txt:
        add_context(source + "#ocr", txt)
        process_text(txt, source + "#ocr", extr)
    
    # Análisis de diagrama
    for it in analyze_diagram(pil_img, source + "#diagram"):
        extr.add(it.kind, it.text, it.source, it.meta)


def process_path(path: str, extr: Extraction):
    if os.path.isdir(path):
        for root, _, files in os.walk(path):
            for f in files:
                process_path(os.path.join(root, f), extr)
        return

    ext = os.path.splitext(path)[1].lower()
    filename = os.path.basename(path)

    try:
        if ext in [".txt", ".md", ".rst"]:
            text = read_txt(path)
            add_context(filename, text)
            process_text(text, filename, extr)

        elif ext in [".pdf"]:
            full_text, page_texts = read_pdf_text(path)
            if page_texts:
                for (pnum, ptext) in page_texts:
                    origin = f"{filename}, pag. {pnum}"
                    add_context(origin, ptext)
                    process_text(ptext, origin, extr)
            else:
                add_context(filename, full_text)
                process_text(full_text, filename, extr)

            if not full_text:
                missing = []
                if pdfplumber is None: missing.append("pdfplumber")
                if fitz is None: missing.append("PyMuPDF")
                if PdfReader is None: missing.append("PyPDF2")
                if missing:
                    extr.add('OBS', f"No se pudo extraer texto de PDF por falta de librerías: {', '.join(missing)}.", filename)

            dfs = extract_pdf_tables(path)
            process_tablelike(dfs, filename, extr)

            if USE_OCR:
                if convert_from_path is None and fitz is None:
                     extr.add('OBS', f"No se pueden analizar diagramas de {filename}. Faltan 'pdf2image' o 'PyMuPDF'.", filename)
                else:
                    print(f"{filename}: Renderizando páginas a imágenes para análisis de diagramas...")
                    full_page_imgs = images_from_pdf(path)
                    
                    if not full_page_imgs:
                         print(f" {filename}: No se pudieron renderizar las páginas a imágenes.")
                    
                    page_text_map = {pnum: ptext for pnum, ptext in page_texts}

                    for (page_idx, pil_img) in full_page_imgs:
                        page_content = page_text_map.get(page_idx, "").lower()
                        is_diagram = False

                        if any(trigger in page_content for trigger in DIAGRAM_TRIGGERS):
                            is_diagram = True
                        elif 50 < len(page_content) < 1500:
                            is_diagram = True
                        elif len(page_content.strip()) < 50:
                            is_diagram = True

                        if is_diagram:
                            print(f" Página {page_idx} detectada como diagrama. Ejecutando análisis de color...")
                            process_image(pil_img, f"{filename}, pag. {page_idx}#diagrama_completo", extr)
                        else:
                            print(f" Página {page_idx} omitida (solo texto).")
            else:
                print(f"OCR desactivado. Omitiendo análisis de diagramas en {filename}.")
                pass

        elif ext in [".docx"]:
            text, tables = read_docx(path)
            add_context(filename, text)
            process_text(text, filename, extr)
            if pd is not None and tables:
                process_tablelike([(pd.DataFrame(t), None) for t in tables], filename, extr)
            try:
                if docx is not None and Image is not None:
                    d = docx.Document(path)
                    for i, rel in enumerate(d.part.rels.values(), 1):
                        if "image" in rel.reltype:
                            data = rel.target_part.blob
                            pil = Image.open(io.BytesIO(data)).convert('RGB')
                            process_image(pil, f"{filename}:img{i}", extr)
            except Exception:
                pass

        elif ext in [".png", ".jpg", ".jpeg", ".tiff", ".bmp"]:
            if Image is not None:
                pil = Image.open(path).convert('RGB')
                process_image(pil, filename, extr)
        else:
            if Image is not None:
                try:
                    pil = Image.open(path).convert('RGB')
                    process_image(pil, filename, extr)
                except Exception:
                    pass
    except Exception as e:
        extr.add('OBS', f"No se pudo procesar {filename}: {e}", filename)


# -------------------- Self tests --------------------

def _self_tests() -> int:
    errors = 0
    global CONTEXT
    CONTEXT = {}
    s = context_section()
    assert "Contexto General" in s, "context_section no crea título"

    add_context("foo.pdf, pag. 1", "Linea A")
    add_context("foo.pdf, pag. 1", "Linea B")
    s2 = context_section()
    try:
        assert "foo.pdf, pag. 1" in s2 and "Linea A" in s2 and "Linea B" in s2, "context_section no incluye contenido"
    except AssertionError as e:
        print("[TEST]", e); errors += 1

    raw = "El sistema generará un\n\npresupuesto preliminar.\nSe enviará."
    norm = normalize_text(raw)
    try:
        assert "generará un" in norm and "presupuesto preliminar." in norm, "normalize_text no repara saltos duros"
    except AssertionError as e:
        print("[TEST]", e); errors += 1

    ex = Extraction()
    process_text("El sistema debe validar el NIF. Si el NIF es inválido, mostrará error.", "demo.txt", ex)
    try:
        assert any("debe" in i.text.lower() for i in ex.rf), "No detecta RF"
        assert any(i.kind == 'DC' for i in ex.dc), "No detecta DC"
    except AssertionError as e:
        print("[TEST]", e); errors += 1

    i = 3; t = "hola"; _tmp = ""; _tmp += f"\n\n[PAGE {i}]\n" + t
    try:
        assert "[PAGE 3]" in _tmp and _tmp.endswith("hola"), "F-string page marker incorrecta"
    except AssertionError as e:
        print("[TEST]", e); errors += 1

    return errors

# -------------------- CLI principal --------------------

def main():
    ap = argparse.ArgumentParser(description="Extrae requisitos y reglas → Markdown (con CONTEXTO GENERAL)")
    ap.add_argument('inputs', nargs='*', help='Rutas de archivos o carpetas de entrada')
    ap.add_argument('-o', '--output', default='salida.md', help='Fichero Markdown de salida')
    ap.add_argument('--tesseract-path', default=None, help='Ruta del binario de tesseract si no está en PATH')
    ap.add_argument('--include-context', action='store_true', help='Incluye la sección de Contexto General en la salida')
    ap.add_argument('--selftest', action='store_true', help='Ejecuta tests de humo y sale')
    args = ap.parse_args()

    if args.selftest:
        failed = _self_tests()
        if failed:
            print(f"Selftests fallidos: {failed}")
            sys.exit(1)
        print("Selftests OK")
        sys.exit(0)

    global USE_OCR
    USE_OCR = False
    if args.tesseract_path:
        if pytesseract is None:
            print("--tesseract-path indicado, pero 'pytesseract' no está disponible; no se realizará OCR.")
        else:
            pytesseract.pytesseract.tesseract_cmd = args.tesseract_path
            USE_OCR = True

    if not args.inputs:
        print("No se proporcionaron entradas. Ejemplo: python req_extractor.py input.pdf -o salida.md")
        sys.exit(2)

    # Crear carpeta de imágenes si no existe
    os.makedirs(OUTPUT_IMG_DIR, exist_ok=True)

    extr = Extraction()

    for p in args.inputs:
        process_path(p, extr)

    conflicts = detect_conflicts(extr.rf + extr.rb)
    for a, b, sim in conflicts:
        extr.add('OBS', f"Posible discrepancia entre: '{a.text}' y '{b.text}' (similitud {sim}%). Revisar.", f"{a.source} vs {b.source}")

    md = extr.to_markdown(include_context = args.include_context)
    with open(args.output, 'w', encoding='utf-8') as f:
        f.write(md)

    print(f"Generado: {args.output}")
    print(f"Imágenes extraídas guardadas en: {os.path.abspath(OUTPUT_IMG_DIR)}")


if __name__ == '__main__':
    main()